# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from .models import CheckTask, CheckResult


def write_docx_report(task: CheckTask, results: list[CheckResult], out_dir: str | Path) -> Path:
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise RuntimeError('缺少 python-docx，请执行：python -m pip install python-docx') from e

    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    path = out / f'{task.name}_{task.id or "new"}.docx'

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    try:
        from docx.oxml.ns import qn
        normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass

    title = doc.add_heading('配置核查报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_table(rows=3, cols=4)
    meta.style = 'Light Grid Accent 1'
    pairs = [
        ('任务名称', task.name), ('设备类型', task.devtype),
        ('设备品牌', task.brand), ('配置文件', task.config_path),
        ('生成时间', task.created_at), ('检查项数量', str(len(results))),
    ]
    for i, (k, v) in enumerate(pairs):
        r = i // 2
        c = (i % 2) * 2
        meta.cell(r, c).text = k
        meta.cell(r, c + 1).text = v or ''

    doc.add_heading('结果汇总', level=1)
    counts = {}
    for r in results:
        counts[r.status] = counts.get(r.status, 0) + 1
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    tbl.rows[0].cells[0].text = '状态'
    tbl.rows[0].cells[1].text = '数量'
    for k, v in counts.items():
        cells = tbl.add_row().cells
        cells[0].text = k
        cells[1].text = str(v)

    doc.add_heading('检查项明细', level=1)
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = 'Light Grid Accent 1'
    headers = ['检查项', '风险', '状态', '证据', '原因', '整改建议']
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    for r in results:
        cells = tbl.add_row().cells
        vals = [r.baseline_name, r.risk, r.status, r.evidence, r.reason, r.suggestion]
        for i, v in enumerate(vals):
            cells[i].text = v or ''
            if i == 2:
                for p in cells[i].paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        if r.status == '未通过':
                            run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
                        elif r.status == '通过':
                            run.font.color.rgb = RGBColor(0x04, 0x78, 0x57)
    doc.save(path)
    return path
